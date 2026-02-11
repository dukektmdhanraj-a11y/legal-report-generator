from flask import Flask, render_template, request, send_file
from docx import Document

app = Flask(__name__)

def replace_all(doc, mapping):
    # Normal paragraphs
    for p in doc.paragraphs:
        for k, v in mapping.items():
            if k in p.text:
                p.text = p.text.replace(k, v)

    # Tables (THIS IS THE FIX)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k, v in mapping.items():
                        if k in p.text:
                            p.text = p.text.replace(k, v)

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        doc = Document("template.docx")

        mapping = {
            "{{DATE}}": request.form.get("DATE",""),
            "{{Branch}}": request.form.get("Branch",""),
            "{{Applicant}}": request.form.get("Applicant",""),
            "{{Co_Applicant}}": request.form.get("Co_Applicant",""),
            "{{TitleHolder}}": request.form.get("TitleHolder",""),
            "{{Loan_Amount}}": request.form.get("Loan_Amount",""),
            "{{AC_No}}": request.form.get("AC_No",""),
            "{{NirnAI_No}}": request.form.get("NirnAI_No",""),
            "{{SY_No}}": request.form.get("SY_No",""),
            "{{DEED_NO}}": request.form.get("DEED_NO",""),

            "{{DISTRICT}}": request.form.get("DISTRICT",""),
            "{{SUB_DISTRICT}}": request.form.get("SUB_DISTRICT",""),
            "{{MANDAL}}": request.form.get("MANDAL",""),
            "{{VILLAGE}}": request.form.get("VILLAGE",""),
            "{{DOOR_NO}}": request.form.get("DOOR_NO",""),
            "{{SURVEY_NO}}": request.form.get("SURVEY_NO",""),
            "{{EXTENT_YARDS}}": request.form.get("EXTENT_YARDS",""),
            "{{EXT_YDS}}": request.form.get("EXT_YDS",""),

            "{{BOUNDARY_NORTH}}": request.form.get("BOUNDARY_NORTH",""),
            "{{BOUNDARY_SOUTH}}": request.form.get("BOUNDARY_SOUTH",""),
            "{{BOUNDARY_EAST}}": request.form.get("BOUNDARY_EAST",""),
            "{{BOUNDARY_WEST}}": request.form.get("BOUNDARY_WEST",""),

            "{{MEASURE_NORTH_EW}}": request.form.get("MEASURE_NORTH_EW",""),
            "{{MEASURE_SOUTH_EW}}": request.form.get("MEASURE_SOUTH_EW",""),
            "{{MEASURE_EAST_NS}}": request.form.get("MEASURE_EAST_NS",""),
            "{{MEASURE_WEST_NS}}": request.form.get("MEASURE_WEST_NS",""),

            "{{ICD_DATE}}": request.form.get("ICD_DATE",""),
            "{{HT_VILLAGE}}": request.form.get("HT_VILLAGE",""),
            "{{HT_DATE}}": request.form.get("HT_DATE",""),
            "{{HT_NAME}}": request.form.get("HT_NAME",""),
            "{{HT_NO}}": request.form.get("HT_NO",""),
            "{{HT_ASSESS_NO}}": request.form.get("HT_ASSESS_NO",""),

            "{{EC_NO}}": request.form.get("EC_NO",""),
            "{{SRO}}": request.form.get("SRO",""),
            "{{EC_DATE}}": request.form.get("EC_DATE",""),
            "{{ECDAY}}": request.form.get("ECDAY",""),
            "{{EC_TIME_PRD}}": request.form.get("EC_TIME_PRD",""),

            "{{NirnAI_EC_No}}": request.form.get("NirnAI_EC_No",""),
            "{{NIRAI_FROM}}": request.form.get("NIRAI_FROM",""),
            "{{NIRAI_TO}}": request.form.get("NIRAI_TO",""),
            "{{Remarks}}": request.form.get("Remarks",""),

            "{{DEED_TYPE}}": request.form.get("DEED_TYPE",""),
            "{{EXECUTOR}}": request.form.get("EXECUTOR",""),
            "{{FAVOUR}}": request.form.get("FAVOUR",""),
            "{{DEED_DATE}}": request.form.get("DEED_DATE",""),
            "{{WORTH}}": request.form.get("WORTH",""),

            "{{PRESENT_OWNER}}": request.form.get("PRESENT_OWNER",""),
        }

        replace_all(doc, mapping)

        output = "LEGAL_REPORT_FINAL.docx"
        doc.save(output)
        return send_file(output, as_attachment=True)

    return render_template("form.html")

if __name__ == "__main__":
    app.run(debug=False)